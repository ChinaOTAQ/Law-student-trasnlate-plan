#!/usr/bin/env python3
"""
docx_helpers.py — reusable Word document helpers for pdf-translate-docx skill.

Provides:
  - Word page-level footnote insertion via direct OOXML (footnotes.xml part)
  - Standard paragraph / heading / run formatting helpers
  - Default typography constants

Typography spec:
  正文: 宋体/Times New Roman, 五号(10.5pt), 1.3× line, 段前0.3行, 段后0.3行
  脚注: 宋体/Times New Roman, 小五号(9pt),  1.1× line, 段前0.1行, 段后0.1行
  脚注序号: 上标（superscript）
  引号: 全角「」，禁止半角 ""
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from lxml import etree

# ── Typography constants ───────────────────────────────────────────────────────

FONT_SIZE_PT   = 10.5        # 五号
FN_SIZE_PT     = 9.0         # 小五号
FONT_NAME      = '宋体'
LATIN_FONT     = 'Times New Roman'
LINE_SPACING   = 1.3         # 正文
FN_LINE_SPACING = 1.1        # 脚注

# 正文: 0.30行 × (10.5 × 1.3) ≈ 4.1 pt
SPACE_BEFORE_PT = round(FONT_SIZE_PT * LINE_SPACING * 0.30, 1)
SPACE_AFTER_PT  = round(FONT_SIZE_PT * LINE_SPACING * 0.30, 1)

# 脚注: 0.10行 × (9 × 1.1) ≈ 1.0 pt
FN_SPACE_PT = round(FN_SIZE_PT * FN_LINE_SPACING * 0.10, 1)  # ≈ 1.0 pt

# ── Footnote OOXML constants ───────────────────────────────────────────────────

_W_NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_FN_URI = '/word/footnotes.xml'
_FN_CT  = ('application/vnd.openxmlformats-officedocument'
           '.wordprocessingml.footnotes+xml')
_FN_RT  = ('http://schemas.openxmlformats.org/officeDocument/2006'
           '/relationships/footnotes')

_FN_XML_SEED = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
             xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1">
    <w:p><w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p><w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>
</w:footnotes>"""


def _get_or_create_footnotes_part(doc):
    """Return (footnotes_part, footnotes_root_lxml_element)."""
    try:
        fn_part = doc.part.part_related_by(_FN_RT)
        fn_root = etree.fromstring(fn_part._blob)
    except KeyError:
        blob = _FN_XML_SEED.encode('utf-8')
        fn_part = Part(PackURI(_FN_URI), _FN_CT, blob, doc.part.package)
        doc.part.relate_to(fn_part, _FN_RT)
        fn_root = etree.fromstring(blob)
    return fn_part, fn_root


def add_footnote(paragraph, fn_id: int, fn_text: str, doc: Document):
    """
    Add a Word page-level footnote (not endnote).

    - Inserts a superscript w:footnoteReference run into *paragraph*.
    - Appends the matching w:footnote node to the document's footnotes part.
    - Footnote paragraph: 小五号(9pt) 宋体/Times New Roman, 1.1× line,
      段前/段后 0.1行, footnote number as superscript.

    Parameters
    ----------
    paragraph : docx Paragraph
    fn_id     : int or str — unique footnote key; must be unique per doc
    fn_text   : str — full footnote body text
    doc       : docx Document
    """
    fn_part, fn_root = _get_or_create_footnotes_part(doc)
    W = _W_NS

    fn_node = etree.SubElement(fn_root, f'{{{W}}}footnote')
    fn_node.set(f'{{{W}}}id', str(fn_id))

    p_fn = etree.SubElement(fn_node, f'{{{W}}}p')
    pPr  = etree.SubElement(p_fn, f'{{{W}}}pPr')

    # paragraph style
    pSty = etree.SubElement(pPr, f'{{{W}}}pStyle')
    pSty.set(f'{{{W}}}val', 'FootnoteText')

    # paragraph spacing: 1.1× line, 段前/段后 ≈ 1pt (0.1行 × 9pt × 1.1)
    fn_sp_twips = str(int(round(FN_SPACE_PT * 20)))   # twips
    fn_line     = str(int(round(FN_LINE_SPACING * 240)))  # 264 for 1.1×
    spacing = etree.SubElement(pPr, f'{{{W}}}spacing')
    spacing.set(f'{{{W}}}before',   fn_sp_twips)
    spacing.set(f'{{{W}}}after',    fn_sp_twips)
    spacing.set(f'{{{W}}}line',     fn_line)
    spacing.set(f'{{{W}}}lineRule', 'auto')

    # superscript footnote number run
    r_num  = etree.SubElement(p_fn, f'{{{W}}}r')
    rPr_n  = etree.SubElement(r_num, f'{{{W}}}rPr')
    rSty_n = etree.SubElement(rPr_n, f'{{{W}}}rStyle')
    rSty_n.set(f'{{{W}}}val', 'FootnoteReference')
    # force superscript on footnote number
    vertAlign = etree.SubElement(rPr_n, f'{{{W}}}vertAlign')
    vertAlign.set(f'{{{W}}}val', 'superscript')
    etree.SubElement(r_num, f'{{{W}}}footnoteRef')

    # footnote text run (小五号=9pt, 宋体/Times New Roman)
    fn_sz = str(int(FN_SIZE_PT * 2))   # half-points: 9pt → '18'
    r_txt  = etree.SubElement(p_fn, f'{{{W}}}r')
    rPr_t  = etree.SubElement(r_txt, f'{{{W}}}rPr')
    for tag in ('sz', 'szCs'):
        el = etree.SubElement(rPr_t, f'{{{W}}}{tag}')
        el.set(f'{{{W}}}val', fn_sz)
    rF = etree.SubElement(rPr_t, f'{{{W}}}rFonts')
    rF.set(f'{{{W}}}eastAsia', FONT_NAME)
    rF.set(f'{{{W}}}ascii',    LATIN_FONT)
    rF.set(f'{{{W}}}hAnsi',    LATIN_FONT)
    t = etree.SubElement(r_txt, f'{{{W}}}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = ' ' + fn_text

    fn_part._blob = etree.tostring(fn_root, xml_declaration=True,
                                   encoding='UTF-8', standalone=True)

    # ── inline superscript reference mark in body paragraph ──────────────────
    run = paragraph.add_run()
    r   = run._r
    rPr = OxmlElement('w:rPr')
    rs  = OxmlElement('w:rStyle')
    rs.set(qn('w:val'), 'FootnoteReference')
    rPr.append(rs)
    # ensure superscript in body too
    va = OxmlElement('w:vertAlign')
    va.set(qn('w:val'), 'superscript')
    rPr.append(va)
    r.append(rPr)
    ref = OxmlElement('w:footnoteReference')
    ref.set(qn('w:id'), str(fn_id))
    r.append(ref)


# ── Run / paragraph helpers ────────────────────────────────────────────────────

def _set_run_font(run, size_pt=FONT_SIZE_PT, bold=False, italic=False):
    """Set font (宋体 + Times New Roman), size, style on a run."""
    run.font.name   = FONT_NAME
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    rPr    = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:ascii'),    LATIN_FONT)
    rFonts.set(qn('w:hAnsi'),    LATIN_FONT)


def fmt_para(para, size_pt=FONT_SIZE_PT, bold=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_pt=21):
    """Apply standard body spacing, line-spacing, and indent to a paragraph."""
    pf = para.paragraph_format
    pf.space_before      = Pt(SPACE_BEFORE_PT)
    pf.space_after       = Pt(SPACE_AFTER_PT)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = LINE_SPACING
    if first_line_pt and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pf.first_line_indent = Pt(first_line_pt)
    if align:
        para.alignment = align
    for run in para.runs:
        _set_run_font(run, size_pt, bold)


def add_body_para(doc, text, bold=False, first_line_pt=21,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add a plain body paragraph, return it."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=bold)
    fmt_para(p, bold=bold, first_line_pt=first_line_pt, align=align)
    return p


def add_heading(doc, text, level=1):
    """Add a heading; level 1 = centered, levels 2–3 = left."""
    p    = doc.add_paragraph()
    run  = p.add_run(text)
    size = FONT_SIZE_PT + (4 - level) * 1.5   # h1≈13.5  h2≈12  h3≈10.5
    _set_run_font(run, size_pt=size, bold=True)
    fmt_para(p, size_pt=size, bold=True,
             align=(WD_ALIGN_PARAGRAPH.CENTER if level == 1
                    else WD_ALIGN_PARAGRAPH.LEFT),
             first_line_pt=0)
    return p


def set_page_margins(doc, top_cm=2.54, bottom_cm=2.54,
                     left_cm=3.17, right_cm=3.17):
    """Set page margins (cm) on the first section."""
    from docx.shared import Cm
    s = doc.sections[0]
    s.top_margin    = Cm(top_cm)
    s.bottom_margin = Cm(bottom_cm)
    s.left_margin   = Cm(left_cm)
    s.right_margin  = Cm(right_cm)


if __name__ == '__main__':
    print('docx_helpers.py — OK')
    print(f'  正文: {FONT_NAME}/{LATIN_FONT} {FONT_SIZE_PT}pt '
          f'×{LINE_SPACING} 段前/后{SPACE_BEFORE_PT}pt')
    print(f'  脚注: {FONT_NAME}/{LATIN_FONT} {FN_SIZE_PT}pt '
          f'×{FN_LINE_SPACING} 段前/后{FN_SPACE_PT}pt')

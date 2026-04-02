#!/usr/bin/env python3
# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx", "lxml"]
# ///
"""
Generic docx builder: reads _structure.txt + _translation.txt + _data.py → _cn.docx

This eliminates the need for per-chapter _main.py files, saving massive tokens.
Agents only need to output a _translation.txt file (one translated block per section,
separated by === markers matching the structure file).

Usage:
    uv run build_docx_from_translation.py \
        --structure ch1_structure.txt \
        --translation ch1_translation.txt \
        --data ch1_data.py \
        --output ch1_cn.docx \
        --title "Hadley v. Baxendale：法律工业化进程中的一项研究" \
        --author "理查德·丹泽格"

_translation.txt format:
    Each block from _structure.txt gets a corresponding translated block.
    Blocks are separated by a line starting with ===BLOCK=== or ---BLOCK---
    TITLE blocks → translated heading text
    TEXT blocks → translated paragraph text with {{FN:N}} for footnote markers

    Example:
        ===TITLE===
        一、引言
        ---TEXT---
        在每年从美国法学院毕业的数以千计的学生中{{FN:1}}，几乎所有人都被要求阅读……
        ---TEXT---
        然而，尽管该案被异常广泛地阅读……

    {{FN:N}} markers are replaced with actual footnote references from _data.py.
    {{FN:*}} for author footnotes.
"""

import argparse
import importlib.util
import os
import re
import sys

SKILL_SCRIPTS = '/Users/zhiyuanqian/.lawvable/skills/pdf-translate-docx/scripts'
sys.path.insert(0, SKILL_SCRIPTS)

from docx_helpers import (
    add_footnote, add_heading, fmt_para, _set_run_font,
    set_page_margins, FONT_SIZE_PT
)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def load_footnotes(data_path):
    """Dynamically load FOOTNOTES dict from a _data.py file."""
    spec = importlib.util.spec_from_file_location("data_module", data_path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod.FOOTNOTES


def parse_translation(trans_path):
    """Parse _translation.txt into list of (type, text) tuples."""
    blocks = []
    current_type = None
    current_lines = []

    with open(trans_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.rstrip('\n')
            if line.startswith('===TITLE===') or line.startswith('=== TITLE'):
                if current_type is not None:
                    blocks.append((current_type, '\n'.join(current_lines).strip()))
                current_type = 'title'
                current_lines = []
            elif line.startswith('---TEXT---') or line.startswith('--- TEXT'):
                if current_type is not None:
                    blocks.append((current_type, '\n'.join(current_lines).strip()))
                current_type = 'text'
                current_lines = []
            elif line.startswith('---LIST---') or line.startswith('--- LIST'):
                if current_type is not None:
                    blocks.append((current_type, '\n'.join(current_lines).strip()))
                current_type = 'text'  # treat list as text
                current_lines = []
            elif line.startswith('---QUOTE---') or line.startswith('--- QUOTE'):
                if current_type is not None:
                    blocks.append((current_type, '\n'.join(current_lines).strip()))
                current_type = 'quote'
                current_lines = []
            else:
                current_lines.append(line)

        # Last block
        if current_type is not None:
            blocks.append((current_type, '\n'.join(current_lines).strip()))

    # Filter out empty blocks
    return [(t, txt) for t, txt in blocks if txt]


def build_paragraph(doc, text, footnotes, first_line_pt=21,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    bold=False, size_pt=FONT_SIZE_PT):
    """Build a paragraph with inline {{FN:N}} markers resolved to real footnotes."""
    p = doc.add_paragraph()
    fmt_para(p, size_pt=size_pt, bold=bold, align=align,
             first_line_pt=first_line_pt)

    # Split text by {{FN:...}} markers
    pattern = r'\{\{FN:([^}]+)\}\}'
    parts = re.split(pattern, text)

    # parts alternates: text, fn_id, text, fn_id, ...
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # Text segment
            if part:
                run = p.add_run(part)
                _set_run_font(run, size_pt=size_pt, bold=bold)
        else:
            # Footnote ID
            fn_id = part.strip()
            # Convert to int if numeric
            try:
                fn_key = int(fn_id)
            except ValueError:
                fn_key = fn_id

            if fn_key in footnotes:
                add_footnote(p, fn_key, footnotes[fn_key], doc)
            else:
                # Fallback: add as superscript text
                run = p.add_run(f'[{fn_id}]')
                _set_run_font(run, size_pt=size_pt, bold=bold)
                run.font.superscript = True

    return p


def build_docx(translation_blocks, footnotes, output_path,
               title=None, author=None):
    """Build the complete docx from parsed translation blocks."""
    doc = Document()
    set_page_margins(doc)

    # Optional title
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        _set_run_font(run, size_pt=16, bold=True)
        fmt_para(p, size_pt=16, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, first_line_pt=0)

    # Optional author
    if author:
        p = doc.add_paragraph()
        run = p.add_run(author)
        _set_run_font(run, size_pt=12, bold=False)
        fmt_para(p, size_pt=12, bold=False,
                 align=WD_ALIGN_PARAGRAPH.CENTER, first_line_pt=0)
        # Add author footnote if exists
        if '*' in footnotes:
            add_footnote(p, '*', footnotes['*'], doc)

    heading_count = 0
    for block_type, text in translation_blocks:
        if block_type == 'title':
            heading_count += 1
            # First 1-2 titles might be chapter title/author — use level 1
            # Subsequent titles are section headings — use level 2
            level = 1 if heading_count <= 1 else 2
            add_heading(doc, text, level=level)

        elif block_type == 'text':
            build_paragraph(doc, text, footnotes)

        elif block_type == 'quote':
            build_paragraph(doc, text, footnotes,
                            first_line_pt=42,
                            align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.save(output_path)
    print(f'Saved: {output_path} ({os.path.getsize(output_path) // 1024} KB)')


def main():
    parser = argparse.ArgumentParser(
        description='Build docx from structure + translation + footnotes')
    parser.add_argument('--structure', help='_structure.txt (optional, for reference)')
    parser.add_argument('--translation', required=True,
                        help='_translation.txt with translated blocks')
    parser.add_argument('--data', required=True,
                        help='_data.py with FOOTNOTES dict')
    parser.add_argument('--output', required=True,
                        help='Output .docx path')
    parser.add_argument('--title', default=None,
                        help='Chapter title (Chinese)')
    parser.add_argument('--author', default=None,
                        help='Author name (Chinese)')

    args = parser.parse_args()

    footnotes = load_footnotes(args.data)
    blocks = parse_translation(args.translation)
    build_docx(blocks, footnotes, args.output,
               title=args.title, author=args.author)


if __name__ == '__main__':
    main()
